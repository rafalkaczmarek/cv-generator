"""Generate the final CV DOCX file from a TailoredCV.

Two responsibilities:
* Ensure built-in Word templates with Jinja2 placeholders exist (and accept
  extra user-dropped `.docx` files in the templates directory).
* Render a chosen template with docxtpl against a TailoredCV and save to /output.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate

from cv_generator.config import get_settings
from cv_generator.models import TailoredCV

_DEFAULT_TEMPLATE_NAME = "cv_template.docx"

_BUILTIN_LABELS: dict[str, tuple[str, str]] = {
    "cv_template.docx": ("Klasyczny", "Ciemnoniebieskie nagłówki, układ standardowy."),
    "cv_modern.docx": ("Nowoczesny", "Wyśrodkowany nagłówek, akcent teal, więcej powietrza."),
    "cv_compact.docx": ("Kompaktowy", "Mniejsza czcionka — więcej treści na jednej stronie."),
}


@dataclass(frozen=True)
class TemplateInfo:
    """A selectable CV Word template."""

    path: Path
    label: str
    description: str = ""

    @property
    def id(self) -> str:
        return self.path.name


def ensure_default_template(template_dir: Path | None = None) -> Path:
    """Create the classic Word template if it doesn't exist. Return its path."""
    return _ensure_template_file(
        template_dir,
        _DEFAULT_TEMPLATE_NAME,
        style="classic",
    )


def ensure_builtin_templates(template_dir: Path | None = None) -> list[Path]:
    """Create all built-in templates that are missing. Return their paths."""
    settings = get_settings()
    template_dir = template_dir or settings.app_templates_dir
    template_dir.mkdir(parents=True, exist_ok=True)
    return [
        _ensure_template_file(template_dir, "cv_template.docx", style="classic"),
        _ensure_template_file(template_dir, "cv_modern.docx", style="modern"),
        _ensure_template_file(template_dir, "cv_compact.docx", style="compact"),
    ]


def list_templates(template_dir: Path | None = None) -> list[TemplateInfo]:
    """Return available templates (built-ins first, then any extra `.docx` files)."""
    settings = get_settings()
    template_dir = template_dir or settings.app_templates_dir
    ensure_builtin_templates(template_dir)

    builtin_names = list(_BUILTIN_LABELS)
    found: dict[str, Path] = {
        p.name: p for p in sorted(template_dir.glob("*.docx")) if p.is_file()
    }

    result: list[TemplateInfo] = []
    for name in builtin_names:
        path = found.pop(name, None)
        if path is None:
            continue
        label, description = _BUILTIN_LABELS[name]
        result.append(TemplateInfo(path=path, label=label, description=description))

    for path in sorted(found.values(), key=lambda p: p.name):
        result.append(
            TemplateInfo(
                path=path,
                label=_humanize_stem(path.stem),
                description="Szablon użytkownika.",
            )
        )
    return result


def resolve_template(
    template_id: str | None = None,
    *,
    template_dir: Path | None = None,
) -> Path:
    """Resolve a template filename (or None → default classic) to an absolute path."""
    templates = list_templates(template_dir)
    if not templates:
        return ensure_default_template(template_dir)
    if template_id is None:
        return templates[0].path
    for info in templates:
        if info.id == template_id:
            return info.path
    raise FileNotFoundError(f"CV template not found: {template_id}")


def render_cv(
    cv: TailoredCV,
    *,
    template_path: Path | None = None,
    template_id: str | None = None,
    output_dir: Path | None = None,
    filename: str | None = None,
) -> Path:
    """Render the CV to a .docx file and return its path."""
    settings = get_settings()
    if template_path is None:
        template_path = resolve_template(template_id)
    output_dir = output_dir or settings.app_output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name_slug = "_".join(cv.full_name.lower().split())
        filename = f"cv_{name_slug}_{stamp}.docx"

    output_path = output_dir / filename
    doc = DocxTemplate(str(template_path))
    doc.render({"cv": cv.model_dump()})
    doc.save(str(output_path))
    return output_path


def _ensure_template_file(
    template_dir: Path | None,
    filename: str,
    *,
    style: str,
) -> Path:
    settings = get_settings()
    template_dir = template_dir or settings.app_templates_dir
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = template_dir / filename
    if template_path.exists():
        return template_path
    _build_template(template_path, style=style)
    return template_path


def _humanize_stem(stem: str) -> str:
    return stem.replace("_", " ").replace("-", " ").strip().title() or stem


def _build_template(template_path: Path, *, style: str) -> None:
    if style == "modern":
        _build_modern_template(template_path)
    elif style == "compact":
        _build_compact_template(template_path)
    else:
        _build_classic_template(template_path)


def _build_classic_template(template_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    accent = RGBColor(0x1F, 0x3A, 0x5F)

    name = doc.add_paragraph()
    name_run = name.add_run("{{ cv.full_name }}")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = accent

    headline = doc.add_paragraph("{{ cv.headline }}")
    headline.runs[0].italic = True
    headline.runs[0].font.size = Pt(12)

    _add_contact_line(doc, size=10)
    _add_cv_body(doc, accent=accent, section_size=13, meta_size=10)
    doc.save(template_path)


def _build_modern_template(template_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    accent = RGBColor(0x0D, 0x73, 0x6B)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run("{{ cv.full_name }}")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = accent

    headline = doc.add_paragraph("{{ cv.headline }}")
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.runs[0].font.size = Pt(12)
    headline.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    contact = _add_contact_line(doc, size=9)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_cv_body(doc, accent=accent, section_size=12, meta_size=9, section_uppercase=False)
    doc.save(template_path)


def _build_compact_template(template_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    accent = RGBColor(0x2C, 0x2C, 0x2C)

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(48)
        section.right_margin = Pt(48)

    name = doc.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    name_run = name.add_run("{{ cv.full_name }}")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = accent

    headline = doc.add_paragraph("{{ cv.headline }}")
    headline.paragraph_format.space_after = Pt(2)
    headline.runs[0].italic = True
    headline.runs[0].font.size = Pt(10)

    _add_contact_line(doc, size=8)
    _add_cv_body(doc, accent=accent, section_size=10, meta_size=8, tight=True)
    doc.save(template_path)


def _add_contact_line(doc: Document, *, size: int):
    contact = doc.add_paragraph(
        "{{ cv.email }}{% if cv.phone %} | {{ cv.phone }}{% endif %}"
        "{% if cv.location %} | {{ cv.location }}{% endif %}"
        "{% if cv.linkedin_url %} | {{ cv.linkedin_url }}{% endif %}"
        "{% if cv.github_url %} | {{ cv.github_url }}{% endif %}"
    )
    contact.runs[0].font.size = Pt(size)
    return contact


def _add_cv_body(
    doc: Document,
    *,
    accent: RGBColor,
    section_size: int,
    meta_size: int,
    section_uppercase: bool = True,
    tight: bool = False,
) -> None:
    _add_section_title(
        doc,
        "Profil",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    summary = doc.add_paragraph("{{ cv.summary }}")
    if tight:
        summary.paragraph_format.space_after = Pt(4)

    _add_section_title(
        doc,
        "Doświadczenie",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    doc.add_paragraph("{%p for exp in cv.experiences %}")
    p = doc.add_paragraph()
    if tight:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
    r = p.add_run("{{ exp.title }} — {{ exp.company }}")
    r.bold = True
    meta = doc.add_paragraph(
        "{{ exp.date_range }}{% if exp.location %} | {{ exp.location }}{% endif %}"
    )
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(meta_size)
    if tight:
        meta.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("{%p for bullet in exp.bullets %}")
    bullet = doc.add_paragraph("{{ bullet }}", style="List Bullet")
    if tight:
        bullet.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endfor %}")

    _add_section_title(
        doc,
        "Wykształcenie",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    doc.add_paragraph("{%p for line in cv.education_lines %}")
    edu = doc.add_paragraph("{{ line }}", style="List Bullet")
    if tight:
        edu.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("{%p endfor %}")

    _add_section_title(
        doc,
        "Umiejętności",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    doc.add_paragraph("{{ cv.skills | join(', ') }}")

    _add_section_title(
        doc,
        "Języki",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    doc.add_paragraph("{{ cv.languages | join(', ') }}")

    _add_section_title(
        doc,
        "Certyfikaty",
        accent=accent,
        size=section_size,
        uppercase=section_uppercase,
        tight=tight,
    )
    doc.add_paragraph("{%p for cert in cv.certifications %}")
    cert = doc.add_paragraph("{{ cert }}", style="List Bullet")
    if tight:
        cert.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("{%p endfor %}")


def _add_section_title(
    doc: Document,
    text: str,
    *,
    accent: RGBColor,
    size: int,
    uppercase: bool = True,
    tight: bool = False,
) -> None:
    p = doc.add_paragraph()
    if tight:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    label = text.upper() if uppercase else text
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = accent
