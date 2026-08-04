"""Tests for DOCX template listing and CV rendering."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from cv_generator.services.docx_generator import (
    ensure_builtin_templates,
    ensure_default_template,
    list_templates,
    render_cv,
    resolve_template,
)


def test_default_template_is_created(tmp_path: Path) -> None:
    template = ensure_default_template(template_dir=tmp_path)
    assert template.exists()
    assert template.suffix == ".docx"
    assert template.stat().st_size > 0


def test_builtin_templates_are_created(tmp_path: Path) -> None:
    paths = ensure_builtin_templates(template_dir=tmp_path)
    assert len(paths) == 3
    assert {p.name for p in paths} == {
        "cv_template.docx",
        "cv_modern.docx",
        "cv_compact.docx",
    }
    assert all(p.exists() and p.stat().st_size > 0 for p in paths)


def test_list_templates_includes_builtins_and_custom(tmp_path: Path) -> None:
    custom = tmp_path / "moj_szablon.docx"
    custom.write_bytes(ensure_default_template(template_dir=tmp_path).read_bytes())

    templates = list_templates(template_dir=tmp_path)
    ids = [t.id for t in templates]
    assert ids[:3] == ["cv_template.docx", "cv_modern.docx", "cv_compact.docx"]
    assert "moj_szablon.docx" in ids
    custom_info = next(t for t in templates if t.id == "moj_szablon.docx")
    assert custom_info.label == "Moj Szablon"


def test_resolve_template_by_id(tmp_path: Path) -> None:
    ensure_builtin_templates(template_dir=tmp_path)
    path = resolve_template("cv_modern.docx", template_dir=tmp_path)
    assert path.name == "cv_modern.docx"


def test_resolve_template_unknown_raises(tmp_path: Path) -> None:
    ensure_builtin_templates(template_dir=tmp_path)
    with pytest.raises(FileNotFoundError, match="missing.docx"):
        resolve_template("missing.docx", template_dir=tmp_path)


def test_render_cv_produces_docx(sample_tailored_cv, tmp_path: Path) -> None:
    template = ensure_default_template(template_dir=tmp_path / "templates")
    output_path = render_cv(
        sample_tailored_cv,
        template_path=template,
        output_dir=tmp_path / "out",
        filename="result.docx",
    )
    assert output_path.exists()
    assert output_path.name == "result.docx"
    assert output_path.stat().st_size > 0


def test_render_cv_with_template_id(sample_tailored_cv, tmp_path: Path) -> None:
    ensure_builtin_templates(template_dir=tmp_path / "templates")
    output_path = render_cv(
        sample_tailored_cv,
        template_id="cv_compact.docx",
        output_dir=tmp_path / "out",
        filename="compact.docx",
    )
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_render_cv_uses_default_filename(sample_tailored_cv, tmp_path: Path) -> None:
    template = ensure_default_template(template_dir=tmp_path / "templates")
    output_path = render_cv(
        sample_tailored_cv,
        template_path=template,
        output_dir=tmp_path / "out",
    )
    assert output_path.name.startswith("cv_jan_kowalski_en_")
    assert output_path.suffix == ".docx"


def test_ensure_default_template_reuses_existing(tmp_path: Path) -> None:
    first = ensure_default_template(template_dir=tmp_path)
    first.write_bytes(b"custom-template")
    second = ensure_default_template(template_dir=tmp_path)
    assert second == first
    assert second.read_bytes() == b"custom-template"


def test_render_cv_includes_courses_section(sample_tailored_cv, tmp_path: Path) -> None:
    template = ensure_default_template(template_dir=tmp_path / "templates")
    output_path = render_cv(
        sample_tailored_cv,
        template_path=template,
        output_dir=tmp_path / "out",
        filename="with_courses.docx",
        language="en",
    )
    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "COURSES" in text
    assert "Kubernetes Fundamentals" in text
    skills_pos = text.index("SKILLS")
    courses_pos = text.index("COURSES")
    languages_pos = text.index("LANGUAGES")
    assert skills_pos < courses_pos < languages_pos


def test_render_cv_uses_polish_section_labels(sample_tailored_cv, tmp_path: Path) -> None:
    template = ensure_default_template(template_dir=tmp_path / "templates")
    pl_cv = sample_tailored_cv.model_copy(update={"language": "pl"})
    output_path = render_cv(
        pl_cv,
        template_path=template,
        output_dir=tmp_path / "out",
        filename="pl.docx",
    )
    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "KURSY" in text
    assert "UMIEJĘTNOŚCI" in text
    assert "JĘZYKI" in text
    assert "COURSES" not in text